import json
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# For pdf conversion
from docx2pdf import convert

def create_resume_docx(data, output_path):
    """
    Creates a Word document from a structured JSON resume.

    Args:
        data (dict): The JSON data containing resume information.
        output_path (str): The path to save the .docx file.
    """
    document = Document()

    # Set Margins for the entire document
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    # Define some styles for cleaner formatting
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(11)

    # Helper function to create a new paragraph with a specific style
    def add_heading(text, level=1):
        return document.add_heading(text, level=level)

    # Helper function to set list style (bullet points or numbered)
    def add_list_item(text, level=0, bullet=True):
        if bullet:
            p = document.add_paragraph(text, style='List Bullet')
        else:
            p = document.add_paragraph(text, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25 * level)
        return p

    # --- Header (Name and Contact Info) ---
    p = document.add_paragraph()
    p_run = p.add_run(data['Name'])
    p_run.bold = True
    p_run.font.size = Pt(15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    # Assuming you have this info from another source
    contact_info = "Phone: (123) 456-7890 | Email: john.doe@email.com"
    p = document.add_paragraph(contact_info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Add a horizontal line separator
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # p._element.get_or_add_pPr().get_or_add_pBdr().get_or_add_bottom().set(qn('w:val'), 'single')

    # --- Sections ---
    # Education Section
    h = document.add_heading("Education", level=2)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(0)
    for degree, degree_info_dict in data['Education'].items():
        p = add_list_item(f"", bullet=True)
        p.add_run(f"{degree}").bold = True
        p.add_run(f", {degree_info_dict['Institution']}, GPA: {degree_info_dict['GPA']}")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Skills Section
    h = document.add_heading("Skills", level=2)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(0)
    for domain, skills_list in data['Skills'].items():
        skills_str = ", ".join(skills_list)
        p = add_list_item(f"", bullet=True)
        p.add_run(f"{domain}").bold = True
        p.add_run(f": {skills_str}")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    # Experience Section
    h = document.add_heading("Exerience", level=2)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(0)
    for company, details in data['Experience'].items():
        p = add_list_item(f"", bullet=False)
        p.add_run(f"{details['Position']}, {company}, {details['Start']} - {details['End']}").bold = True
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for bullet in details['Experience']:
            p2 = add_list_item(bullet, bullet=True, level=1)
            p2.paragraph_format.left_indent = Inches(0.50)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Projects Section
    h = document.add_heading("Projects", level=2)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(0)
    for project, description in data['Projects'].items():
        p = add_list_item(f"", bullet=True)
        p.add_run(f"{project}: ").bold = True
        p.add_run(f"{description}")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Research Publication Section
    if 'Research Publications' in data:
        h = document.add_heading("Research Publications", level=2)
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(0)
        for article_name, details in data['Research Publications'].items():
            p = add_list_item(f"{details['last name']}, {details['first initial']}. \"{article_name}\" {details['journal name']}, {details['volume number']}, {details['date']}", bullet=True)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save Docx file
    document.save(output_path)
    print(f"✅ Resume successfully created at {output_path}")

def create_resume_pdf(input_docx_path, output_path):
    print("Creating PDF")
    convert(input_docx_path, output_path)
    print('✅ Resume successfully created at ', output_path)